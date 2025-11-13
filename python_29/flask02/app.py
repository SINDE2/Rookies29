import os
from flask import Flask, render_template, request, redirect, url_for, send_from_directory
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# â­ í•œê¸€ í°íŠ¸ ë“±ë¡ (ìœˆë„ìš° ê¸°ì¤€ ë§‘ì€ ê³ ë”•)
pdfmetrics.registerFont(
    TTFont('MalgunGothic', r'C:\Windows\Fonts\malgun.ttf')
)


app = Flask(__name__)

# ìˆ˜ë£Œì¦ íŒŒì¼ ì €ì¥ í´ë”
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CERT_DIR = os.path.join(BASE_DIR, "certificates")
os.makedirs(CERT_DIR, exist_ok=True)


@app.route("/", methods=["GET"])
def form_page():
    # ì´ë¦„, ê³¼ì •, ë‚ ì§œ ì…ë ¥ ë°›ëŠ” í˜ì´ì§€
    return render_template("form.html")


@app.route("/generate", methods=["POST"])
def generate_certificate():
    name = request.form.get("name")
    course = request.form.get("course")
    date = request.form.get("date")  # 2025-11-13 ì´ëŸ° í˜•ì‹

    if not (name and course and date):
        return "ëª¨ë“  í•­ëª©ì„ ì…ë ¥í•˜ì„¸ìš”.", 400

    # íŒŒì¼ ì´ë¦„(í•œê¸€/ê³µë°± ë¬¸ì œ ë°©ì§€ë¥¼ ìœ„í•´ ê°„ë‹¨íˆ ì²˜ë¦¬)
    safe_name = name.replace(" ", "_")
    safe_course = course.replace(" ", "_")
    safe_date = date.replace("-", "")
    base_filename = f"{safe_name}_{safe_course}_{safe_date}"

    docx_filename = f"{base_filename}.docx"
    pdf_filename = f"{base_filename}.pdf"

    docx_path = os.path.join(CERT_DIR, docx_filename)
    pdf_path = os.path.join(CERT_DIR, pdf_filename)

    # 1) DOCX ìˆ˜ë£Œì¦ ìƒì„±
    create_docx_certificate(docx_path, name, course, date)

    # 2) PDF ìˆ˜ë£Œì¦ ìƒì„±
    create_pdf_certificate(pdf_path, name, course, date)

    # ê²°ê³¼ í˜ì´ì§€ë¡œ ì´ë™ (PDF ë‹¤ìš´ë¡œë“œ ë§í¬ ì•ˆë‚´)
    return render_template(
        "result.html",
        name=name,
        course=course,
        date=date,
        pdf_filename=pdf_filename,
        docx_filename=docx_filename,
    )


@app.route("/download/pdf/<filename>")
def download_pdf(filename):
    # PDF íŒŒì¼ ë‹¤ìš´ë¡œë“œ
    return send_from_directory(
        CERT_DIR,
        filename,
        as_attachment=True,
        mimetype="application/pdf"
    )


@app.route("/download/docx/<filename>")
def download_docx(filename):
    # DOCX íŒŒì¼ ë‹¤ìš´ë¡œë“œ (ì˜µì…˜)
    return send_from_directory(
        CERT_DIR,
        filename,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def create_docx_certificate(path, name, course, date):
    """python-docxë¡œ ìˆ˜ë£Œì¦ docx ìƒì„±"""
    document = Document()

    document.add_heading("ìˆ˜ ë£Œ ì¦", level=0)

    p = document.add_paragraph()
    p.add_run(f"{name} ë‹˜ì€ ").bold = True
    p.add_run(f"ã€{course}ã€ ê³¼ì •ì„ ì„±ì‹¤íˆ ì´ìˆ˜í•˜ì˜€ìŒì„ í™•ì¸í•©ë‹ˆë‹¤.\n")

    document.add_paragraph("")
    document.add_paragraph(f"ìˆ˜ë£Œì¼ì: {date}")
    document.add_paragraph("ê¸°ê´€ëª…: â—‹â—‹êµìœ¡ì„¼í„°")

    document.add_paragraph("")
    document.add_paragraph("ìœ„ì™€ ê°™ì´ ìˆ˜ë£Œì¦ì„ ë°œê¸‰í•©ë‹ˆë‹¤.")

    document.save(path)


def create_pdf_certificate(path, name, course, date):
    """reportlabìœ¼ë¡œ ìˆ˜ë£Œì¦ pdf ìƒì„± (í•œê¸€ í°íŠ¸ ì‚¬ìš©)"""
    c = canvas.Canvas(path, pagesize=A4)
    width, height = A4

    # ğŸ”¹ ì œëª©
    c.setFont("MalgunGothic", 28)
    c.drawCentredString(width / 2, height - 120, "ìˆ˜ ë£Œ ì¦")

    # ğŸ”¹ ë³¸ë¬¸
    c.setFont("MalgunGothic", 14)
    text_y = height - 200
    c.drawString(80, text_y, f"ì„±    ëª… : {name}")
    text_y -= 30
    c.drawString(80, text_y, f"ê³¼    ì • : {course}")
    text_y -= 30
    c.drawString(80, text_y, f"ìˆ˜ë£Œì¼ì : {date}")

    text_y -= 50
    c.setFont("MalgunGothic", 12)
    c.drawString(
        80,
        text_y,
        "ìœ„ ì‚¬ëŒì€ ìœ„ ê³¼ì •ì„ ì„±ì‹¤íˆ ì´ìˆ˜í•˜ì˜€ìœ¼ë¯€ë¡œ ì´ ìˆ˜ë£Œì¦ì„ ìˆ˜ì—¬í•©ë‹ˆë‹¤."
    )

    text_y -= 80
    c.setFont("MalgunGothic", 14)
    c.drawRightString(width - 80, text_y, "â—‹â—‹êµìœ¡ì„¼í„°ì¥  (ì¸)")

    c.showPage()
    c.save()



if __name__ == "__main__":
    app.run(debug=True)
